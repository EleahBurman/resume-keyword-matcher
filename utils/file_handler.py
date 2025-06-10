"""
File Handler Module for Resume Keyword Matcher

This module handles extracting text from different file formats:
- PDF files using PyPDF2
- DOCX files using python-docx
- TXT files (plain text)

Author: Resume Keyword Matcher
"""

import os
from pathlib import Path
import PyPDF2
from docx import Document


class FileHandler:
    """
    Handles file operations and text extraction from various file formats
    """
    
    def __init__(self):
        """Initialize the FileHandler"""
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
    
    def extract_text(self, file_path):
        """
        Extract text from a file based on its extension
        
        Args:
            file_path (str): Path to the file to extract text from
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
            Exception: For other file processing errors
        """
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        file_extension = Path(file_path).suffix.lower()
        
        # Check if file type is supported
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            # Route to appropriate extraction method based on file type
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file extension: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_path}: {str(e)}")
    
    def _extract_from_pdf(self, file_path):
        """
        Extract text from PDF file using PyPDF2
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # Add page text to overall content
                    if page_text:
                        text_content += page_text + "\n"
                
                # Clean up the extracted text
                text_content = self._clean_extracted_text(text_content)
                
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        
        return text_content
    
    def _extract_from_docx(self, file_path):
        """
        Extract text from DOCX file using python-docx
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        try:
            # Load the document
            doc = Document(file_path)
            
            # Extract text from each paragraph
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_content += paragraph.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            # Clean up the extracted text
            text_content = self._clean_extracted_text(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
        
        return text_content
    
    def _extract_from_txt(self, file_path):
        """
        Extract text from plain text file
        
        Args:
            file_path (str): Path to TXT file
            
        Returns:
            str: File content
        """
        try:
            # Try different encodings in case of special characters
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return self._clean_extracted_text(content)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, raise an error
            raise Exception("Could not decode file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def _clean_extracted_text(self, text):
        """
        Clean and normalize extracted text
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Strip whitespace from each line
            cleaned_line = line.strip()
            
            # Skip empty lines and lines with just special characters
            if cleaned_line and len(cleaned_line) > 1:
                cleaned_lines.append(cleaned_line)
        
        # Join lines with single spaces, but preserve paragraph breaks
        cleaned_text = ' '.join(cleaned_lines)
        
        # Remove multiple consecutive spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def validate_file(self, file_path):
        """
        Validate if a file can be processed
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: Validation result with status and message
        """
        result = {
            'valid': False,
            'message': '',
            'file_type': None,
            'file_size': 0
        }
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                result['message'] = 'File does not exist'
                return result
            
            # Get file info
            file_extension = Path(file_path).suffix.lower()
            file_size = os.path.getsize(file_path)
            
            result['file_type'] = file_extension
            result['file_size'] = file_size
            
            # Check file extension
            if file_extension not in self.supported_extensions:
                result['message'] = f'Unsupported file type: {file_extension}'
                return result
            
            # Check file size (16MB limit)
            max_size = 16 * 1024 * 1024  # 16MB in bytes
            if file_size > max_size:
                result['message'] = f'File too large: {file_size / (1024*1024):.1f}MB (max 16MB)'
                return result
            
            # Check if file is empty
            if file_size == 0:
                result['message'] = 'File is empty'
                return result
            
            result['valid'] = True
            result['message'] = 'File is valid'
            
        except Exception as e:
            result['message'] = f'Error validating file: {str(e)}'
        
        return result
    
    def get_supported_extensions(self):
        """
        Get list of supported file extensions
        
        Returns:
            list: List of supported extensions
        """
        return list(self.supported_extensions)


# Test function to verify the module works
def test_file_handler():
    """Test function for FileHandler - run this to verify everything works"""
    print("Testing FileHandler...")
    
    handler = FileHandler()
    
    # Test supported extensions
    print(f"Supported extensions: {handler.get_supported_extensions()}")
    
    # Test with sample text
    test_text = """
    John Doe
    Software Developer
    
    Experience:
    - Python development
    - Flask web applications
    - Database management
    
    Skills: Python, JavaScript, SQL, Git
    """
    
    # Create a temporary test file
    import tempfile
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write(test_text)
        temp_file = f.name
    
    try:
        # Test extraction
        extracted = handler.extract_text(temp_file)
        print(f"Extracted text length: {len(extracted)} characters")
        print(f"First 100 characters: {extracted[:100]}...")
        
        # Test validation
        validation = handler.validate_file(temp_file)
        print(f"Validation result: {validation}")
        
        print("✅ FileHandler test passed!")
        
    except Exception as e:
        print(f"❌ FileHandler test failed: {e}")
    
    finally:
        # Clean up
        os.unlink(temp_file)


if __name__ == "__main__":
    # Run test when file is executed directly
    test_file_handler()